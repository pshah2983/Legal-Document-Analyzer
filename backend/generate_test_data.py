import os
import sys
from docx import Document

# Define the output directory for test documents
OUTPUT_DIR = '/Users/parvashah/Desktop/DS Projects/Legal Document Analyzer/test_documents'

# Standard boilerplate templates for building rich, legal documents programmatically
CLAUSE_TEMPLATES = {
    "confidentiality_low": (
        "Confidentiality & Non-Disclosure.\n"
        "Each party agrees to maintain the strict confidentiality of all proprietary or non-public information "
        "disclosed by the other party. This obligation shall not apply to information that: (a) is or becomes "
        "publicly known through no breach of this Agreement; (b) was already in the receiving party's possession "
        "prior to disclosure; or (c) is independently developed without reference to the disclosing party's "
        "Confidential Information. These obligations shall survive for a period of two (2) years following the "
        "termination of this Agreement."
    ),
    "confidentiality_high": (
        "Unilateral Absolute Confidentiality Obligations.\n"
        "The Receiving Party shall maintain absolute confidentiality of all information disclosed by the Disclosing "
        "Party under this Agreement. Under no circumstances, including court subpoenas or legally binding orders, "
        "shall the Receiving Party disclose any such information without the express prior written consent of the "
        "Disclosing Party. The Receiving Party agrees that its confidentiality obligations under this Agreement "
        "shall survive in perpetuity and shall extend to all affiliates, subcontractors, and related entities."
    ),
    "indemnity_low": (
        "Mutual Indemnification.\n"
        "Each party (the 'Indemnifying Party') agrees to indemnify, defend, and hold harmless the other party "
        "from and against any third-party claims, liabilities, or losses arising directly out of the Indemnifying "
        "Party's gross negligence, willful misconduct, or material breach of representation under this Agreement."
    ),
    "indemnity_high": (
        "Unilateral Broad Indemnification Obligations.\n"
        "The Contractor agrees to indemnify, defend, and hold harmless the Client, its officers, directors, "
        "employees, and agents from and against any and all claims, damages, losses, liabilities, costs, and expenses "
        "(including legal fees and court costs) of any kind, whether direct, indirect, special, or consequential, "
        "arising out of or in connection with the services provided, even if such claims are caused in part by "
        "the negligence or joint misconduct of the Client."
    ),
    "liability_low": (
        "Limitation of Liability (Reciprocal Cap).\n"
        "Except for breaches of confidentiality or indemnification obligations, neither party shall be liable "
        "to the other for any indirect, special, incidental, or consequential damages. In no event shall "
        "either party's total aggregate liability arising under this Agreement exceed the total amounts paid "
        "or payable to the Contractor in the twelve (12) month period preceding the event giving rise to liability."
    ),
    "liability_high": (
        "Unilateral Limitation of Liability (Severe Cap).\n"
        "The Client's total maximum aggregate liability for any and all breaches, claims, or damages arising "
        "under this Agreement, whether in contract, tort, or otherwise, shall be strictly limited to the "
        "sum of One Hundred Dollars ($100.00). The Contractor waives any and all rights to claim consequential, "
        "punitive, or exemplary damages, or lost profits of any nature whatsoever."
    ),
    "termination_low": (
        "Termination for Convenience (Mutual).\n"
        "Either party may terminate this Agreement at any time, for any reason or no reason, by providing "
        "thirty (30) days prior written notice to the other party. Upon termination, all outstanding payments "
        "for services rendered up to the date of termination shall become immediately due and payable."
    ),
    "termination_high": (
        "Unilateral Termination at Will.\n"
        "The Client reserves the absolute right to terminate this Agreement immediately at any time, with or "
        "without cause, by providing oral or written notice to the Contractor. Upon receipt of such notice, "
        "the Contractor shall immediately cease all performance. The Client shall have no obligation to pay "
        "for any services, materials, or work-in-progress completed after the moment of termination."
    ),
    "noncompete_high": (
        "Restrictive Covenant & Non-Compete Agreement.\n"
        "For a period of five (5) years following the termination or expiration of this Agreement, the "
        "Employee shall not, directly or indirectly, engage in, operate, manage, consult for, or be employed by "
        "any business entity anywhere in the world that competes, or intends to compete, with the business "
        "activities or products of the Company as conducted at any time during the Employee's employment."
    ),
    "arbitration_med": (
        "Mandatory Binding Arbitration & Waiver of Jury Trial.\n"
        "Any dispute, controversy, or claim arising out of or relating to this contract, including its formation "
        "or breach, shall be settled solely by binding arbitration administered by the American Arbitration "
        "Association (AAA) in accordance with its Commercial Arbitration Rules. The parties waive their right "
        "to a trial by jury or to participate in any class action lawsuit."
    ),
    "governing_low": (
        "Governing Law.\n"
        "This Agreement shall be governed by, and construed in accordance with, the laws of the State of New York, "
        "without regard to its conflict of law principles. Any dispute arising hereunder shall be subject to the "
        "exclusive jurisdiction of the state and federal courts located in New York County, New York."
    ),
    "governing_high": (
        "Governing Law & Foreign Jurisdiction.\n"
        "This Agreement and all disputes arising hereunder shall be governed solely by the laws of the Republic "
        "of Singapore. The parties agree that any legal action, suit, or proceeding must be filed and adjudicated "
        "exclusively in the courts of Singapore, and the Contractor hereby waives any objection based on forum "
        "non-conveniens."
    )
}

# Define 25 distinct, realistic test agreements commonly used by small businesses
CONTRACTS = [
    {
        "filename": "01_mutual_nda_standard.docx",
        "title": "Mutual Non-Disclosure Agreement",
        "intro": "This Mutual NDA is entered into between AuroraTech Solutions and Nebula Labs for discussions regarding a strategic partnership.",
        "clauses": ["confidentiality_low", "governing_low"]
    },
    {
        "filename": "02_independent_contractor_agreement.docx",
        "title": "Independent Contractor Services Agreement",
        "intro": "This agreement details terms for engineering services provided by Jane Doe to Apex Digital Brands Corp.",
        "clauses": ["confidentiality_low", "indemnity_high", "liability_low", "termination_low", "governing_low"]
    },
    {
        "filename": "03_commercial_office_lease.docx",
        "title": "Commercial Real Estate Lease Agreement",
        "intro": "This Commercial Lease binds Brick & Mortar Holdings (Landlord) and Sweet Crumbs Bakery (Tenant) for the premises at 404 Market St.",
        "clauses": ["indemnity_high", "liability_high", "termination_high", "arbitration_med", "governing_low"]
    },
    {
        "filename": "04_software_license_agreement.docx",
        "title": "End User Software License Agreement (EULA)",
        "intro": "This EULA governs the download and use of the CloudScribe SaaS editor developed by ByteForge Inc.",
        "clauses": ["confidentiality_high", "liability_high", "arbitration_med", "governing_high"]
    },
    {
        "filename": "05_employment_agreement_strict.docx",
        "title": "Executive Employment & Non-Compete Agreement",
        "intro": "This Executive Employment Agreement is made between CyberGuard Systems and Alexander Mercer, Chief Technology Officer.",
        "clauses": ["confidentiality_high", "noncompete_high", "termination_high", "arbitration_med", "governing_low"]
    },
    {
        "filename": "06_master_services_agreement.docx",
        "title": "Master Services Agreement (MSA)",
        "intro": "This MSA establishes the core framework under which PeakMarketing LLC will execute media campaigns for Horizon Fitness Group.",
        "clauses": ["confidentiality_low", "indemnity_low", "liability_low", "termination_low", "governing_low"]
    },
    {
        "filename": "07_statement_of_work_web_dev.docx",
        "title": "Statement of Work (SOW) - Web Development",
        "intro": "This SOW defines deliverables for the Shopify storefront redevelopment project signed by CodeCrafters and Blossom Florist.",
        "clauses": ["confidentiality_low", "liability_low", "termination_low", "governing_low"]
    },
    {
        "filename": "08_vendor_supply_contract.docx",
        "title": "Standard Vendor Raw Material Supply Agreement",
        "intro": "This Supply Agreement governs purchase orders of premium organic flour between WheatFields Millers and Breadwinners Co.",
        "clauses": ["indemnity_low", "liability_high", "termination_low", "governing_low"]
    },
    {
        "filename": "09_unilateral_nda_investor.docx",
        "title": "Unilateral Investor Non-Disclosure Agreement",
        "intro": "This Unilateral NDA protects proprietary AI research presented by VentureScale Fund to Silicon Valley Biotech Inc.",
        "clauses": ["confidentiality_high", "governing_low"]
    },
    {
        "filename": "10_equipment_rental_agreement.docx",
        "title": "Heavy Construction Equipment Rental Agreement",
        "intro": "Governs the temporary rental of high-capacity earth excavators from IronFleet Rentals to Cornerstone Builders.",
        "clauses": ["indemnity_high", "liability_high", "arbitration_med", "governing_low"]
    },
    {
        "filename": "11_saas_terms_of_service.docx",
        "title": "SaaS Platform Terms of Service",
        "intro": "Governs access to the TaskFlow enterprise collaborative project board for all registered commercial accounts.",
        "clauses": ["confidentiality_low", "liability_high", "arbitration_med", "governing_high"]
    },
    {
        "filename": "12_marketing_agency_agreement.docx",
        "title": "Exclusive Marketing & Social Media Agency Agreement",
        "intro": "Binds Prime Influencers Inc. and Velvet Glow Cosmetics for exclusive promotional campaigns on social networks.",
        "clauses": ["confidentiality_low", "indemnity_high", "termination_low", "governing_low"]
    },
    {
        "filename": "13_referral_partner_agreement.docx",
        "title": "Business Referral & Affiliate Partner Agreement",
        "intro": "Establishes referral fees and commissions payable by LeadGenerator Co. to GrowthHacks Consultants.",
        "clauses": ["confidentiality_low", "liability_low", "termination_low", "governing_low"]
    },
    {
        "filename": "14_intellectual_property_assignment.docx",
        "title": "Proprietary Information & Intellectual Property Assignment",
        "intro": "Ensures all codes, assets, and concepts built during contracting belong strictly to Vanguard Fintech Systems.",
        "clauses": ["confidentiality_high", "indemnity_high", "governing_low"]
    },
    {
        "filename": "15_franchise_agreement_excerpt.docx",
        "title": "Franchise Agreement Excerpt & Operating Terms",
        "intro": "Binds QuickBurger Corp (Franchisor) and local operations franchisee partnership Capital Burgers LLC.",
        "clauses": ["indemnity_high", "noncompete_high", "arbitration_med", "governing_high"]
    },
    {
        "filename": "16_consulting_agreement_template.docx",
        "title": "Strategic Business Consulting Agreement",
        "intro": "Outlines corporate advisory service scopes provided by Apex Advisory Partners to Summit Logistics Inc.",
        "clauses": ["confidentiality_low", "liability_low", "termination_low", "governing_low"]
    },
    {
        "filename": "17_subcontractor_agreement.docx",
        "title": "Subcontractor General Services Agreement",
        "intro": "Governs HVAC mechanical install tasks outsourced by HVAC-Prime Inc. to Thermal-Flow Mechanicals.",
        "clauses": ["confidentiality_low", "indemnity_high", "liability_low", "governing_low"]
    },
    {
        "filename": "18_event_space_rental_contract.docx",
        "title": "Commercial Event Venue Rental Agreement",
        "intro": "Rental of the Grand Ballroom at Majestic Plaza for the annual TechCon Innovate gala hosted by Startups-United.",
        "clauses": ["indemnity_high", "liability_high", "termination_high", "governing_low"]
    },
    {
        "filename": "19_catering_services_contract.docx",
        "title": "Standard Corporate Catering Services Agreement",
        "intro": "Binds GourmetBites Caterers and EnterpriseHub Co. for weekly staff lunch provisions and holiday events.",
        "clauses": ["indemnity_low", "liability_low", "termination_low", "governing_low"]
    },
    {
        "filename": "20_beta_testing_agreement.docx",
        "title": "Beta Software Tester Agreement & Feedback Licence",
        "intro": "Governs early closed-beta tests of the SoundFlow audio synthesiser built by WaveForm Labs.",
        "clauses": ["confidentiality_high", "liability_high", "governing_low"]
    },
    {
        "filename": "21_freelance_graphic_design.docx",
        "title": "Freelance Design Services & Asset Delivery Agreement",
        "intro": "Governs vector brand identity assets designed by Lily Creative for Sparkle Dental Practice.",
        "clauses": ["confidentiality_low", "termination_low", "governing_low"]
    },
    {
        "filename": "22_photography_service_contract.docx",
        "title": "Commercial Studio Photography Service Agreement",
        "intro": "Covers product photoshoot schedules commissioned by EcoWear Apparel to Focus Lens Studios.",
        "clauses": ["liability_low", "termination_low", "governing_low"]
    },
    {
        "filename": "23_co_marketing_agreement.docx",
        "title": "Joint Co-Marketing & Cross-Promotion Agreement",
        "intro": "Joint advertising agreement between TravelBound Agency and StayGlobal Hotel Alliance.",
        "clauses": ["confidentiality_low", "indemnity_low", "governing_low"]
    },
    {
        "filename": "24_content_licensing_agreement.docx",
        "title": "Digital Media & Video Content Licensing Agreement",
        "intro": "Licences high-definition drone travel footage from SkyView Media to Wanderlust TV Network.",
        "clauses": ["confidentiality_low", "indemnity_high", "liability_low", "governing_low"]
    },
    {
        "filename": "25_noncompete_employee_only.docx",
        "title": "Standard Employee Non-Compete Covenant",
        "intro": "A single Restrictive Covenant document binding sales associates at Direct-Logistics Supply Corp.",
        "clauses": ["noncompete_high", "arbitration_med", "governing_low"]
    }
]

def generate_documents():
    """Generates all 25 agreements as distinct Microsoft Word files in output folder."""
    if not os.path.exists(OUTPUT_DIR):
        print(f"Creating output directory: {OUTPUT_DIR}")
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        
    print(f"Starting programmatical generation of {len(CONTRACTS)} contract files...")
    
    for idx, c in enumerate(CONTRACTS):
        doc = Document()
        
        # Add Header Styling
        doc.add_heading(c["title"], level=1)
        doc.add_paragraph("==============================================================================")
        doc.add_paragraph("LEGAL AGREEMENT - COMMERCIAL RECORD")
        doc.add_paragraph("==============================================================================")
        
        # Add Introduction
        doc.add_heading("1. Parties & Recitals", level=2)
        doc.add_paragraph(c["intro"])
        doc.add_paragraph(
            "Whereas, the parties hereto wish to formalize their business relationship in accordance "
            "with the terms and conditions outlined below. Now, therefore, in consideration of the mutual "
            "covenants contained herein, the parties agree as follows:"
        )
        
        # Loop through mapped clauses and write them
        doc.add_heading("2. Specific Operating Covenants", level=2)
        for i, c_key in enumerate(c["clauses"]):
            clause_text = CLAUSE_TEMPLATES[c_key]
            # Paragraph formatting
            p = doc.add_paragraph()
            # Title bolding
            lines = clause_text.split('\n')
            p.add_run(f"Section 2.{i+1} — {lines[0]}\n").bold = True
            p.add_run(lines[1])
            
        # Add Signature Boilerplates
        doc.add_heading("3. Execution & Signatures", level=2)
        doc.add_paragraph(
            "IN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the date "
            "first written above."
        )
        p_sig = doc.add_paragraph()
        p_sig.add_run("Party A Representative: _______________________    Date: ______________\n\n")
        p_sig.add_run("Party B Representative: _______________________    Date: ______________")
        
        # Save document
        file_path = os.path.join(OUTPUT_DIR, c["filename"])
        doc.save(file_path)
        print(f"[{idx+1}/{len(CONTRACTS)}] Generated: {c['filename']}")

    print("\nSuccess! All 25 commercial agreements have been successfully generated.")
    print(f"Output Directory: file://{OUTPUT_DIR}")

if __name__ == '__main__':
    generate_documents()
